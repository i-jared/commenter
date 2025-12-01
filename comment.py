#!/usr/bin/env python3
import argparse
import json
import os
import re
import sys
from typing import List, Union, Dict, Any

# --- DOCX imports ---
from docx import Document

# --- PDF imports ---
import fitz  # PyMuPDF

# --- OpenAI import ---
from openai import OpenAI

def extract_text(file_path: str) -> str:
    """Extracts text content from .docx, .pdf, or text files."""
    if not file_path or not os.path.exists(file_path):
        return ""
    
    ext = os.path.splitext(file_path)[1].lower()
    
    try:
        if ext == ".docx":
            doc = Document(file_path)
            return "\n".join([p.text for p in doc.paragraphs])
        elif ext == ".pdf":
            doc = fitz.open(file_path)
            text = ""
            for page in doc:
                text += page.get_text()
            return text
        else:
            # Assume text file
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
    except Exception as e:
        print(f"Error extracting text from {file_path}: {e}")
        return ""

def generate_annotations(doc_path: str, rubric_path: str, assign_kb_path: str, general_kb_path: str) -> List[Dict[str, Any]]:
    """
    Generates annotations using OpenAI API based on the input document and auxiliary files.
    """
    client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))

    # Extract text from all documents
    doc_text = extract_text(doc_path)
    rubric_text = extract_text(rubric_path)
    assign_kb_text = extract_text(assign_kb_path)
    general_kb_text = extract_text(general_kb_path)

    system_prompt = """
    You are a helpful teaching assistant. Your goal is to grade a paper and provide feedback.
    You will be provided with the student's paper, a rubric, an assignment knowledge base, and a general knowledge base.
    
    Your output must be a valid JSON object containing a list of annotations. 
    The JSON structure should be exactly as follows:
    [
      {
        "target": {
          "text": "exact text from the paper to attach the comment to",
          "occurrence": "first" 
        },
        "comment": {
          "text": "The feedback comment",
          "author": "AI Grader"
        }
      }
    ]
    
    Ensure the "text" in "target" matches a specific sentence or phrase in the student's paper EXACTLY so it can be located.
    """

    user_content = f"""
    ## Student Paper (Input):
    {doc_text}
    
    ## Rubric:
    {rubric_text}
    
    ## Assignment Knowledge Base:
    {assign_kb_text}
    
    ## General Knowledge Base:
    {general_kb_text}
    
    Grade this paper based on the provided context.
    """

    try:
        completion = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_content}
            ],
            response_format={"type": "json_object"}
        )
        
        content = completion.choices[0].message.content
        
        # Parse the JSON response
        # The model might return { "annotations": [...] } or just [...]
        data = json.loads(content)
        
        if isinstance(data, list):
            return data
        elif isinstance(data, dict):
            # Check if it's wrapped in a key like 'annotations'
            for key in data:
                if isinstance(data[key], list):
                    return data[key]
            # If it's a single annotation object
            return [data]
        else:
            return []

    except Exception as e:
        raise RuntimeError(f"OpenAI API Error: {e}")

def load_annotations(json_path: str) -> List[Dict[str, Any]]:
    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    if isinstance(data, dict):
        return [data]
    if isinstance(data, list):
        return data
    raise ValueError("JSON must be an object or a list of objects.")


# ==========================================================
# DOCX COMMENTING
# ==========================================================

def add_comment_to_paragraph(doc: Document, paragraph, comment_text: str, author: str = "Author"):
    """
    Attach a Word comment to an entire paragraph.
    """
    # Use the native add_comment method available in python-docx >= 1.2.0
    # It handles creating the comments part if needed.
    if not paragraph.runs:
        return

    doc.add_comment(paragraph.runs, comment_text, author=author)


def match_text_in_paragraph(
    paragraph_text: str,
    target_text: str,
    match_type: str,
    case_sensitive: bool,
    whole_word: bool,
) -> bool:
    """
    Return True if paragraph_text contains target_text according to the options.
    This is intentionally simple and operates at paragraph granularity.
    """
    if not case_sensitive:
        paragraph_text_cmp = paragraph_text.lower()
        target_text_cmp = target_text.lower()
    else:
        paragraph_text_cmp = paragraph_text
        target_text_cmp = target_text

    if match_type == "regex":
        flags = 0
        if not case_sensitive:
            flags |= re.IGNORECASE

        pattern = target_text
        if whole_word:
            pattern = r"\b" + pattern + r"\b"

        return re.search(pattern, paragraph_text, flags) is not None
    else:
        # exact / literal substring match
        if whole_word:
            # crude whole word: regex word boundary
            pattern = r"\b" + re.escape(target_text_cmp) + r"\b"
            return re.search(pattern, paragraph_text_cmp) is not None
        return target_text_cmp in paragraph_text_cmp


def annotate_docx(input_path: str, output_path: str, annotations: List[Dict[str, Any]]):
    doc = Document(input_path)

    for ann in annotations:
        target = ann.get("target", {})
        comment_spec = ann.get("comment", {})
        mode = target.get("mode", "text")

        if mode != "text":
            # For brevity, we only implement text targeting here.
            continue

        text = target.get("text", "")
        if not text:
            continue

        match_type = target.get("match_type", "exact")
        case_sensitive = target.get("case_sensitive", False)
        whole_word = target.get("whole_word", True)
        occurrence = target.get("occurrence", "first")

        comment_text = comment_spec.get("text", "")
        author = comment_spec.get("author", "Reviewer")

        # We'll count matches over paragraphs in reading order
        match_count = 0

        for paragraph in doc.paragraphs:
            para_text = paragraph.text or ""
            if not para_text.strip():
                continue

            if not match_text_in_paragraph(
                para_text, text, match_type, case_sensitive, whole_word
            ):
                continue

            # We have a match in this paragraph
            match_count += 1

            if occurrence == "all":
                add_comment_to_paragraph(doc, paragraph, comment_text, author)
                continue

            # occurrence can be "first" or a specific integer
            if occurrence == "first":
                if match_count == 1:
                    add_comment_to_paragraph(doc, paragraph, comment_text, author)
                    break
            else:
                try:
                    occ_num = int(occurrence)
                except Exception:
                    occ_num = 1
                if match_count == occ_num:
                    add_comment_to_paragraph(doc, paragraph, comment_text, author)
                    break

    doc.save(output_path)


# ==========================================================
# PDF ANNOTATION
# ==========================================================

def annotate_pdf(input_path: str, output_path: str, annotations: List[Dict[str, Any]]):
    doc = fitz.open(input_path)

    for ann in annotations:
        target = ann.get("target", {})
        comment_spec = ann.get("comment", {})
        mode = target.get("mode", "text")

        comment_text = comment_spec.get("text", "")
        author = comment_spec.get("author", "Reviewer")

        if mode == "position":
            # Position-based mode: expect page + bbox
            pdf_pos = target.get("pdf", {})
            page_index = pdf_pos.get("page", 1) - 1  # user 1-based
            bbox = pdf_pos.get("bbox")  # [x1, y1, x2, y2]

            if bbox is None or not (0 <= page_index < len(doc)):
                continue

            page = doc[page_index]
            rect = fitz.Rect(*bbox)
            annot = page.add_text_annot(rect.tl, f"{author}: {comment_text}")
            annot.update()
            continue

        # Text-based mode
        text = target.get("text", "")
        if not text:
            continue

        occurrence = target.get("occurrence", "first")
        # match_type, case_sensitive, whole_word are not fully handled here;
        # PyMuPDF search_for is literal and reasonably good for basic use.

        matches_global: List[tuple] = []  # (page_index, rect)

        for page_index in range(len(doc)):
            page = doc[page_index]
            rects = page.search_for(text)  # simple literal search
            for r in rects:
                matches_global.append((page_index, r))

        if not matches_global:
            continue

        def apply_comment_to_rects(rects: List[tuple]):
            for p_idx, rect in rects:
                page = doc[p_idx]
                # highlight the text and attach the comment
                highlight = page.add_highlight_annot(rect)
                highlight.set_info(info={"content": f"{author}: {comment_text}"})
                highlight.update()

        if occurrence == "all":
            apply_comment_to_rects(matches_global)
        else:
            # first or nth
            if occurrence == "first":
                idx = 1
            else:
                try:
                    idx = int(occurrence)
                except Exception:
                    idx = 1

            if 1 <= idx <= len(matches_global):
                apply_comment_to_rects([matches_global[idx - 1]])

    doc.save(output_path)
    doc.close()


# ==========================================================
# CLI
# ==========================================================

def main():
    parser = argparse.ArgumentParser(
        description="Add comments/annotations to DOCX/PDF based on a JSON spec."
    )
    parser.add_argument("document_path", help="Path to input .docx or .pdf")
    # Optional now as we might not use it if we were doing LLM via CLI, 
    # but for backwards compatibility we'll keep it or make it optional?
    # For now, let's keep the original CLI structure as is for manual use, 
    # but the GUI uses the library functions directly.
    parser.add_argument("json_path", help="Path to JSON file with annotation specs")
    parser.add_argument(
        "-o",
        "--output",
        help="Output path (defaults to <name>-annotated<ext>)",
    )
    args = parser.parse_args()

    doc_path = args.document_path
    json_path = args.json_path

    if not os.path.isfile(doc_path):
        print(f"Document not found: {doc_path}", file=sys.stderr)
        sys.exit(1)

    if not os.path.isfile(json_path):
        print(f"JSON file not found: {json_path}", file=sys.stderr)
        sys.exit(1)

    annotations = load_annotations(json_path)

    base, ext = os.path.splitext(doc_path)
    if args.output:
        out_path = args.output
    else:
        out_path = f"{base}-annotated{ext}"

    ext_lower = ext.lower()
    if ext_lower == ".docx":
        annotate_docx(doc_path, out_path, annotations)
    elif ext_lower == ".pdf":
        annotate_pdf(doc_path, out_path, annotations)
    else:
        print("Only .docx and .pdf are supported.", file=sys.stderr)
        sys.exit(1)

    print(f"Wrote annotated file to: {out_path}")


if __name__ == "__main__":
    main()
